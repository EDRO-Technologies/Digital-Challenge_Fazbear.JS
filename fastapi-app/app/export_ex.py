#-*- coding: utf-8 -*-

from peewee import *
from pandas import *
from faker import Faker
from datetime import datetime, timedelta
from docx import Document
from reportlab.lib.pagesizes import *
from reportlab. platypus import SimpleDocTemplate, Table, TableStyle
from openpyxl import *
from reportlab.lib import colors




# Игнор ошибки стилей таблиц
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="docx")


# query = Users.select()


fake = Faker()

# Создание структуры таблицы
table_data = [
    {
        'date': fake.date_this_decade(),
        'time': fake.time(),
        'messages': fake.random_int(min=1, max=100),
        'shift_start': fake.date_time_this_year(),
        'shift_end': fake.date_time_this_year() + timedelta(hours=8),
        'visas': fake.random_int(min=0, max=10),
        'notes': fake.text(),
    } for _ in range(10)
]

# Преобразование данных в DataFrame pandas
data = {'Date': [], 'Time': [], 'Messages during shift': [], 'Shift acceptance': [], 'End of shift': [], 'Visas': [], 'Remarks': []}
for row in table_data:
    data['Date'].append(row['date'])
    data['Time'].append(row['time'])
    data['Messages during shift'].append(row['messages'])
    data['Shift acceptance'].append(row['shift_start'])
    data['End of shift'].append(row['shift_end'])
    data['Visas'].append(row['visas'])
    data['Remarks'].append(row['notes'])


df = DataFrame(data)

current_time = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")

excel_file_path = f'Отчет_{current_time}.xlsx'


df.to_excel(excel_file_path, index=False, engine='openpyxl')

wb = load_workbook(excel_file_path)
ws = wb.active

for column in ws.columns:
    max_length = 0
    column = [cell for cell in column]
    for cell in column:
        try:
            if len(str(cell.value)) > max_length:
                max_length = len(cell.value)
        except:
            pass
    adjusted_width = (max_length + 7)
    ws.column_dimensions[column[0].column_letter].width = adjusted_width
    
wb.save(excel_file_path)

print(f'Data has been exported to {excel_file_path}')

# Сохранение данных в Word
word_creation_time = current_time
word_file_path = f'Отчет_{word_creation_time}.docx'
document = Document()

table = document.add_table(rows=1, cols=len(df.columns))
table.autofit = False

table.style = 'TableGrid'

for col_num, col_name in enumerate(df.columns):
    cell = table.cell(0, col_num)
    cell.text = col_name
    cell.paragraphs[0].runs[0].font.name = 'Cambria'
    cell.paragraphs[0].runs[0].font.bold = True
    

for row_num, row_data in enumerate(df.itertuples(), start=1):
    table.add_row()
    for col_num, value in enumerate(row_data[1:], start=0):
        cell = table.cell(row_num, col_num)
        cell.text = str(value)
        
        
document.save(word_file_path)
print(f'Data has been exported to {word_file_path}')



def convert_word_table_to_pdf(word_file, pdf_file):
    
    doc = Document(word_file)

    pdf = SimpleDocTemplate(pdf_file, pagesize=A2)

    font_size = 10


    data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        data.extend(table_data)


    # last_column_widths = [max(len(str(row[-1])) * 3 for row in data)]


    col_widths = [70, 70, 50, 100, 100, 50, 750] 
    # + last_column_widths

    # Создание таблицы в ReportLab
    pdf_table = Table(data, colWidths=col_widths)


    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 15),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])

    pdf_table.setStyle(style)

    pdf_content = [pdf_table]

    pdf.build(pdf_content)


word_file_path = f"Отчет_{word_creation_time}.docx"
pdf_file_path = f"Отчет_{current_time}.pdf"
convert_word_table_to_pdf(word_file_path, pdf_file_path)
print(f'Data has been exported to {pdf_file_path}')

